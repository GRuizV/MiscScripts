'''THIS WAS MADE TO HANDLE THE RONNIE CHENGS'S RUTINE TRANSCRIPT'''
# import re

# # Define your regex pattern
# pattern = re.compile(r'(\d+:\d+) -')

# # Read the file and modify lines
# with open(r'C:\Users\USUARIO\GR\Software Development\Context Managers\Test - Ronny Chieng Rutine Transcript\file.txt', 'r') as file:
    
#     content = file.read()
#     modified_content = re.sub(pattern, r'\1 - ', content)

# # Write modified lines back to the file
# with open(r'C:\Users\USUARIO\GR\Software Development\Context Managers\Test - Ronny Chieng Rutine Transcript\file.txt', 'w') as file:
    # file.write(modified_content)


'THIS WAS MADE TO HANDLE THE GLOSARY NEWLINES'
# import re

# # Define your regex pattern
# pattern = re.compile(r'\n')

# # Read the file and modify lines
# with open(r'C:\Users\USUARIO\GR\Software Development\Context Managers\Glosary.txt', 'r', encoding='utf-8') as file:
       
#     content = file.read()
#     modified_content = re.sub(pattern, r'\n\n', content)

# # Write modified lines back to the file
# with open(r'C:\Users\USUARIO\GR\Software Development\Context Managers\Glosary.txt', 'w', encoding='utf-8') as file:
#     file.write(modified_content)

    

'THIS WAS MADE TO ADD BOLD FORMAT IN WORD FOR THE GLOSARY'
import re
from docx import Document
from docx.shared import Pt

# Define your regex pattern
pattern = re.compile(r'^(.*?):')


# Create a new Word document
doc = Document()


# Function to handle non-printable characters
def clean_text(text):
    return ''.join(c for c in text if c.isprintable())


# Read the text file
with open(r'C:\Users\USUARIO\GR\Software Development\Context Managers\Glosary\Glosary.docx', 'r', encoding='utf-8', errors='ignore') as file:
    lines = file.readlines()


# Process each line
for line in lines:

    line = clean_text(line)
    match = pattern.match(line)

    if match:

        bold_text = match.group(0)
        normal_text = line[len(bold_text):]

        # Add a new paragraph to the document
        p = doc.add_paragraph()

        # Add the bold part
        run = p.add_run(bold_text)
        run.bold = True
        run.font.size = Pt(12)  # Adjust font size as needed

        # Add the rest of the line
        p.add_run(normal_text)

    else:
        doc.add_paragraph(line)


# Save the document
doc.save(r'C:\Users\USUARIO\GR\Software Development\Context Managers\Glosary\Glosary.docx')



